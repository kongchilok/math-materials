# -*- coding: utf-8 -*-
"""
以《IEP第8-9點_初三數學_26-27上學期_融合班通用草稿.docx》為版面母版，
產出「初三數學補救班（課後輔導）」全學年版 IEP 第 8、9 點。

內容來源：抽離教學_數學/_行政表格/課後輔導教學教案_初三數學補救班_全學年.md（31 節）
使用者裁決（2026-09-03）：① IEP 出「全學年合併一份」②「上」「下」學期兩格皆剔。
私隱：本檔為「通用草稿」，不寫任何學生姓名（跟 CLAUDE.md 鐵律 6）。
"""
import copy, shutil
from pathlib import Path
from docx import Document
from docx.table import _Cell
from docx.oxml.ns import qn

BASE = Path(r"G:\我的雲端硬碟\2ndBrain\notebookLM\抽離教學_數學\初三數學")
SRC  = BASE / "IEP第8-9點_初三數學_26-27上學期_融合班通用草稿.docx"
OUT  = BASE / "IEP第8-9點_初三數學補救班_26-27全學年_融合班通用草稿.docx"

# ---------- 低層工具 ----------
def tcs(row):
    return row._tr.findall(qn('w:tc'))

def set_para(p, text):
    """保留段落格式，只換文字：首個 run 承載全文，其餘 run 移除。"""
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._r.getparent().remove(r._r)

def set_cell(cell, lines):
    """把儲存格改成 lines（每行一段），沿用第一段的段落格式與字型。"""
    if isinstance(lines, str):
        lines = [lines]
    paras = cell.paragraphs
    tmpl_p = paras[0]._p
    for p in paras[1:]:
        p._p.getparent().remove(p._p)
    set_para(cell.paragraphs[0], lines[0])
    anchor = cell.paragraphs[0]._p
    for line in lines[1:]:
        new_p = copy.deepcopy(tmpl_p)
        anchor.addnext(new_p)
        anchor = new_p
        from docx.text.paragraph import Paragraph
        set_para(Paragraph(new_p, cell), line)

def cell_of(tbl, ri, ci):
    return _Cell(tcs(tbl.rows[ri])[ci], tbl)

# ---------- 內容 ----------
SERVICE_CONTENT = (
    "透過課後輔導（補救教學），沿用學生已學的講義、練習與工具卡，"
    "以重做基礎題及錯題訂正的方式，鞏固方程、函數、圓與相似三角的基礎技能，"
    "建立分步驟解題與自我檢查的習慣。了解學生的理解程度。"
)
SERVICE_GOALS = [
    "1. 能在提示步驟卡與公式卡下，重做已學單元的基礎題",
    "2. 能說出該單元解題的關鍵步驟與所用公式",
    "3. 能指出自己錯題的成因並完成訂正",
    "4. 能圈出仍未掌握的題目並主動提問",
]
LONG_TERM = [
    "1. 能重做並訂正一元二次方程、分式方程與二元二次方程組的基礎題，說出所用方法與驗根步驟",
    "2. 能重做並訂正二次函數圖像平移，以及圖像與 x 軸交點個數的基礎題",
    "3. 能重做並訂正旋轉、中心對稱與圓的性質、切線及兩圓位置關係的作圖與計算題",
    "4. 能重做並訂正概率、反比例函數與圖形相似的基礎題",
    "5. 能重做並訂正銳角三角函數與解直角三角形的應用題，並辨識投影的兩種類型",
]
# (短期目標, 起迄日期, 評量方式, 輔助措施)  —— 日期按教案 31 節的實際上課月份
SHORT_TERM = [
    ("1-1. 能重做直接開平方法的題目，並按常數項的正負判斷根的個數", "2026/09 – 2026/09", "B", "a11：手順卡"),
    ("1-2. 能重做因式分解法的題目，並指出「兩邊同除以未知數」會漏根", "2026/09 – 2026/09", "B", "a11：手順卡"),
    ("1-3. 能重做判別式的題目，說出三種情況各自對應的根的個數", "2026/09 – 2026/09", "B", "a11：公式卡"),
    ("1-4. 能重做增長率應用題，寫出設方程的步驟", "2026/10 – 2026/10", "B", "a3,a4"),
    ("1-5. 能重做分式方程應用題，並完成驗根步驟", "2026/10 – 2026/10", "B", "a3,a6"),
    ("1-6. 能重做二元二次方程組的代入消元題，並把答案寫成一對數", "2026/10 – 2026/10", "B", "a11：手順卡"),
    ("1-7. 能在找錯題中圈出漏了驗根的位置並改正", "2026/10 – 2026/10", "B", "a3"),
    ("2-1. 能重做二次函數圖像平移的題目，按「先左右、後上下」的次序覆述", "2026/11 – 2026/11", "B", "a11：手順卡"),
    ("2-2. 能重做用判別式判斷圖像與 x 軸交點個數的題目", "2026/11 – 2026/11", "B", "a11：公式卡"),
    ("3-1. 能重做 90°、180° 旋轉的作圖題，每點量度相同角度與距離", "2026/11 – 2026/12", "C", "a3"),
    ("3-2. 能重做旋轉與中心對稱的綜合題，說出 180° 旋轉即中心對稱", "2026/12 – 2026/12", "B", "a3"),
    ("3-3. 能重做垂徑定理的應用題", "2026/12 – 2026/12", "B", "a11：公式卡"),
    ("3-4. 能重做點和圓的位置關係與切線長定理的題目", "2026/12 – 2027/01", "B", "a11：公式卡"),
    ("3-5. 能重做兩圓位置關係的辨識題，說出五種情況", "2027/01 – 2027/01", "B", "a11：對照卡"),
    ("4-1. 能重做隨機事件分類的題目，並說出概率的定義", "2027/01 – 2027/02", "B", "a3"),
    ("4-2. 能重做放回與不放回的對照題，說出兩者結果數不同", "2027/02 – 2027/02", "B", "a11：對照卡"),
    ("4-3. 能重做用頻率估計概率的題目", "2027/02 – 2027/02", "B", "b5"),
    ("4-4. 能重做用列表法求概率的挑戰題，逐一核對有沒有漏數", "2027/03 – 2027/03", "B", "a3,a6"),
    ("4-5. 能重做反比例函數圖像與性質的辨識題", "2027/03 – 2027/03", "C", "a3,a6"),
    ("5-1. 能重做比例線段的交叉相乘題", "2027/03 – 2027/03", "B", "a11：公式卡"),
    ("5-2. 能重做平行線分線段成比例的應用題，核對對應關係有沒有配錯", "2027/04 – 2027/04", "B", "a3"),
    ("5-3. 能重做相似三角形的判定題，先說出用哪一種判定法", "2027/04 – 2027/04", "B", "a11：判定卡"),
    ("5-4. 能完成相似三角形與位似的綜合練習並訂正錯題", "2027/04 – 2027/04", "B", "a3,a6"),
    ("6-1. 能重做正弦、餘弦、正切三個比值的對照計算題", "2027/05 – 2027/05", "B", "a11：公式卡,b5"),
    ("6-2. 能用計算機由三角函數值求角度，說出反函數的按鍵次序", "2027/05 – 2027/05", "C", "b5"),
    ("6-3. 能覆述特殊角記憶表並完成錯題訂正", "2027/05 – 2027/05", "B", "a11：公式卡"),
    ("6-4. 能重做坡度應用題", "2027/06 – 2027/06", "B", "a11：公式卡,b5"),
    ("6-5. 能辨識中心投影與平行投影兩種情況", "2027/06 – 2027/06", "B", "a3"),
    ("6-6. 能運用公式紙完成圓與幾何部分的總複習題", "2027/06 – 2027/06", "B", "a11：公式紙"),
    ("能圈出題目所求的內容。", "2026/09 – 2027/06", "B", ""),
    ("能用底線標記出題目的已知條件。", "2026/09 – 2027/06", "B", ""),
    ("在附件提供下，能寫出題目所涉及的公式。", "2026/09 – 2027/06", "B", "a11：公式卡"),
    ("能在錯題本記錄錯誤成因，並在下一節前完成訂正。", "2026/09 – 2027/06", "D", ""),
]
REMARK = ("本表為課後輔導（補救教學）版本，逢週四第九節（16:15–16:55），"
          "全學年 31 節，服務對象跨初三三個班別；抽離課堂另有獨立 IEP 表。")

# ---------- 建置 ----------
TMP = OUT.with_name("_tmp_" + OUT.name)     # 先寫暫存，避免 Word 開住目標檔時整爛佢
shutil.copyfile(SRC, TMP)
doc = Document(str(TMP))
t8, t9 = doc.tables[0], doc.tables[1]

# ===== 第 8 點：學習輔導欄（R6 起，vMerge restart 在 R6） =====
set_cell(cell_of(t8, 6, 1), SERVICE_CONTENT)
set_cell(cell_of(t8, 6, 2), SERVICE_GOALS)
set_cell(cell_of(t8, 6, 3), "課後輔導(小組教學)")
set_cell(cell_of(t8, 6, 4), "1")                    # 頻率：次
set_cell(cell_of(t8, 6, 6), "1")                    # 頻率：週
set_cell(cell_of(t8, 6, 7), "2026/09 – 2027/06")
set_cell(cell_of(t8, 6, 8), "江志樂(數學資源老師)")
set_cell(cell_of(t8, 6, 9), "本校課室")
set_cell(cell_of(t8, 8, 4), "40")                   # 每次所需分鐘
set_cell(cell_of(t8, 27, 0), REMARK)                # 【備註】下方空行

# ===== 第 9 點表頭 =====
set_cell(cell_of(t9, 1, 1), ["數學科", "(課後輔導)"])

# 學期：兩個 FORMCHECKBOX 都剔（原檔「上」=1、「下」=0）
sem_tc = tcs(t9.rows[1])[5]
for cb in sem_tc.iter(qn('w:checkBox')):
    dflt = cb.find(qn('w:default'))
    if dflt is None:
        dflt = cb.makeelement(qn('w:default'), {}); cb.append(dflt)
    dflt.set(qn('w:val'), "1")
    ck = cb.find(qn('w:checked'))
    if ck is not None:
        cb.remove(ck)

# 長期學習目標：只換 P1–P5，保留「如有需要:」與備注段
lt_cell = cell_of(t9, 3, 0)
for i, txt in enumerate(LONG_TERM, start=1):
    set_para(lt_cell.paragraphs[i], txt)

# 課程調適：只改「其他（…）」的措辭，剔選框原封不動
adapt_cell = cell_of(t9, 4, 0)
_swap = {"小組": "課後", "抽離": "補救輔導", "以提升專注度及課堂參與": "以鞏固基礎、訂正錯題"}
_hit = False
for r in adapt_cell.paragraphs[4].runs:
    if r.text == "其他":
        _hit = True
        continue
    if _hit and r.text in _swap:
        r.text = _swap.pop(r.text)
assert not _swap, f"課程調適未命中的 run：{_swap}"

# ===== 短期學習目標：清走原 R6–R29，改鋪新列 =====
tmpl_mid  = copy.deepcopy(t9.rows[7]._tr)    # 純文字列（無 OMML）
tmpl_last = copy.deepcopy(t9.rows[29]._tr)   # 末列（保留底框線）
tbl = t9._tbl
for tr in [t9.rows[i]._tr for i in range(6, 30)]:
    tbl.remove(tr)

anchor = t9.rows[5]._tr                      # 表頭列
for idx, (goal, dates, mode, aid) in enumerate(SHORT_TERM):
    last = (idx == len(SHORT_TERM) - 1)
    tr = copy.deepcopy(tmpl_last if last else tmpl_mid)
    anchor.addnext(tr); anchor = tr
    row = t9.rows[6 + idx]
    set_cell(cell_of(t9, 6 + idx, 0), goal)
    set_cell(cell_of(t9, 6 + idx, 1), dates)
    set_cell(cell_of(t9, 6 + idx, 2), mode)
    set_cell(cell_of(t9, 6 + idx, 3), aid)
    set_cell(cell_of(t9, 6 + idx, 4), "")

doc.save(str(TMP))
try:
    import os
    os.replace(TMP, OUT)
    print("OK ->", OUT)
except PermissionError:
    print("!! 目標檔被佔用（Word 開住），已保留暫存檔：", TMP)
    print("!! 關閉 Word 後再跑一次本腳本即可原地替換。")
    raise SystemExit(3)
print("短期目標列數：", len(SHORT_TERM))
